"""
Convert dissertation Markdown files to Word (.docx) format.
Handles headings, paragraphs, bold/italic, tables, and bullet points.
"""
import re
import os
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT


def apply_inline_formatting(paragraph, text):
    """Apply bold and italic formatting to inline text."""
    # Pattern to match **bold**, *italic*, and plain text
    pattern = r'(\*\*\*(.+?)\*\*\*|\*\*(.+?)\*\*|\*(.+?)\*|([^*]+))'
    matches = re.finditer(pattern, text)
    for match in matches:
        if match.group(2):  # ***bold italic***
            run = paragraph.add_run(match.group(2))
            run.bold = True
            run.italic = True
        elif match.group(3):  # **bold**
            run = paragraph.add_run(match.group(3))
            run.bold = True
        elif match.group(4):  # *italic*
            run = paragraph.add_run(match.group(4))
            run.italic = True
        elif match.group(5):  # plain text
            paragraph.add_run(match.group(5))


def md_to_docx(md_path, docx_path):
    """Convert a markdown file to a Word document."""
    doc = Document()
    
    # Set default font
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Times New Roman'
    font.size = Pt(12)
    
    # Set heading styles
    for i in range(1, 5):
        heading_style = doc.styles[f'Heading {i}']
        heading_style.font.name = 'Times New Roman'
        heading_style.font.color.rgb = RGBColor(0, 0, 0)
    
    with open(md_path, 'r', encoding='utf-8') as f:
        lines = f.readlines()
    
    i = 0
    while i < len(lines):
        line = lines[i].rstrip('\n')
        
        # Skip empty lines
        if not line.strip():
            i += 1
            continue
        
        # Horizontal rule
        if line.strip() in ('---', '***', '___'):
            # doc.add_paragraph('─' * 50)
            i += 1
            continue
        
        # Headings
        heading_match = re.match(r'^(#{1,4})\s+(.+)', line)
        if heading_match:
            level = len(heading_match.group(1))
            text = heading_match.group(2).strip()
            heading = doc.add_heading(level=level)
            apply_inline_formatting(heading, text)
            i += 1
            continue
        
        # Table detection
        if '|' in line and i + 1 < len(lines) and '|' in lines[i + 1]:
            # Collect all table rows
            table_lines = []
            while i < len(lines) and '|' in lines[i].strip():
                stripped = lines[i].strip()
                # Skip separator rows (|---|---|)
                if re.match(r'^[\|\s\-:]+$', stripped):
                    i += 1
                    continue
                cells = [c.strip() for c in stripped.split('|')]
                cells = [c for c in cells if c != '']  # Remove empty edge cells
                table_lines.append(cells)
                i += 1
            
            if len(table_lines) >= 1:
                num_cols = len(table_lines[0])
                table = doc.add_table(rows=len(table_lines), cols=num_cols)
                table.style = 'Table Grid'
                table.alignment = WD_TABLE_ALIGNMENT.CENTER
                
                for row_idx, row_data in enumerate(table_lines):
                    for col_idx, cell_text in enumerate(row_data):
                        if col_idx < num_cols:
                            cell = table.cell(row_idx, col_idx)
                            cell.text = ''
                            para = cell.paragraphs[0]
                            # Bold the header row
                            clean_text = re.sub(r'\*\*(.+?)\*\*', r'\1', cell_text)
                            run = para.add_run(clean_text)
                            run.font.size = Pt(10)
                            run.font.name = 'Times New Roman'
                            if row_idx == 0:
                                run.bold = True
                
                doc.add_paragraph('')  # spacing after table
            continue
        
        # Bullet points
        bullet_match = re.match(r'^(\s*)[-*]\s+(.+)', line)
        if bullet_match:
            indent_level = len(bullet_match.group(1)) // 2
            text = bullet_match.group(2)
            para = doc.add_paragraph(style='List Bullet')
            para.paragraph_format.left_indent = Inches(0.25 * (indent_level + 1))
            apply_inline_formatting(para, text)
            i += 1
            continue
        
        # Numbered lists
        num_match = re.match(r'^(\s*)\d+\.\s+(.+)', line)
        if num_match:
            text = num_match.group(2)
            para = doc.add_paragraph(style='List Number')
            apply_inline_formatting(para, text)
            i += 1
            continue
        
        # Regular paragraph — collect continuation lines
        para_text = line
        i += 1
        while i < len(lines):
            next_line = lines[i].rstrip('\n')
            if (not next_line.strip() or 
                next_line.startswith('#') or 
                next_line.startswith('|') or
                next_line.startswith('---') or
                re.match(r'^[-*]\s+', next_line) or
                re.match(r'^\d+\.\s+', next_line)):
                break
            para_text += ' ' + next_line
            i += 1
        
        para = doc.add_paragraph()
        para.paragraph_format.space_after = Pt(6)
        para.paragraph_format.line_spacing = 1.5
        apply_inline_formatting(para, para_text)
    
    doc.save(docx_path)
    print(f"  ✅ Created: {docx_path}")


if __name__ == '__main__':
    base_dir = os.path.dirname(os.path.abspath(__file__))
    
    files_to_convert = [
        '00_Front_Matter.md',
        'Chapter_1_Introduction.md',
        'Chapter_2_Literature_Review.md',
        'Chapter_3_Methodology.md',
        'Chapter_4_Implementation.md',
    ]
    
    print("\n  📄 Converting dissertation files to Word (.docx)...\n")
    
    for md_file in files_to_convert:
        md_path = os.path.join(base_dir, md_file)
        if os.path.exists(md_path):
            docx_file = md_file.replace('.md', '.docx')
            docx_path = os.path.join(base_dir, docx_file)
            md_to_docx(md_path, docx_path)
        else:
            print(f"  ⚠️  Skipped (not found): {md_file}")
    
    print("\n  Done! Files are in the dissertation/ folder.\n")
